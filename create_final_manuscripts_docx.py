import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def create_gastroenteritis_manuscript_docx():
    """Create DOCX version of the updated gastroenteritis manuscript with LOS analysis"""

    # Read the updated markdown content
    with open('comprehensive_gastroenteritis_manuscript.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set up styles
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(14)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(12)
    heading2_style.font.bold = True

    # Parse and format the markdown content
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if line.startswith('# '):
            # Title
            p = doc.add_paragraph(line[2:], style='CustomTitle')
            p.paragraph_format.space_after = Pt(20)
        elif line.startswith('## '):
            # Heading 1
            p = doc.add_paragraph(line[3:], style='CustomHeading1')
            p.paragraph_format.space_after = Pt(12)
        elif line.startswith('### '):
            # Heading 2
            p = doc.add_paragraph(line[4:], style='CustomHeading2')
            p.paragraph_format.space_after = Pt(8)
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        elif line.startswith('|') and '|' in line:
            # Table
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1  # Adjust for the while loop increment

            if len(table_lines) > 1:
                # Parse table
                headers = [col.strip() for col in table_lines[0].split('|')[1:-1]]
                data_rows = []

                for table_line in table_lines[2:]:  # Skip header and separator
                    if table_line.strip():
                        row_data = [col.strip() for col in table_line.split('|')[1:-1]]
                        data_rows.append(row_data)

                # Create table
                if data_rows:
                    table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
                    table.style = 'Table Grid'

                    # Add headers
                    hdr_cells = table.rows[0].cells
                    for j, header in enumerate(headers):
                        hdr_cells[j].text = header
                        hdr_cells[j].paragraphs[0].runs[0].bold = True

                    # Add data
                    for row_idx, row_data in enumerate(data_rows):
                        row_cells = table.rows[row_idx + 1].cells
                        for col_idx, cell_data in enumerate(row_data):
                            row_cells[col_idx].text = cell_data

                    doc.add_paragraph()  # Space after table
        elif line.startswith('![') and '](' in line:
            # Image
            alt_text = line.split('](')[0][2:]
            image_path = line.split('](')[1][:-1]
            try:
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(5))
                    doc.add_paragraph(f"Figure: {alt_text}")
                else:
                    doc.add_paragraph(f"[Image not found: {image_path}]")
            except:
                doc.add_paragraph(f"[Image could not be loaded: {image_path}]")
        elif line.startswith('* ') or line.startswith('- '):
            # List item
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line:
            # Regular paragraph
            p = doc.add_paragraph(line)
        else:
            # Empty line
            doc.add_paragraph()

        i += 1

    # Save document
    doc.save('comprehensive_gastroenteritis_manuscript_final.docx')
    print("Gastroenteritis manuscript DOCX created successfully!")

def create_respiratory_manuscript_docx():
    """Create DOCX version of the updated respiratory manuscript with LOS analysis"""

    # Read the updated markdown content
    with open('comprehensive_respiratory_manuscript.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set up styles (same as gastroenteritis)
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(14)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(12)
    heading2_style.font.bold = True

    # Parse and format the markdown content (same logic as gastroenteritis)
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if line.startswith('# '):
            # Title
            p = doc.add_paragraph(line[2:], style='CustomTitle')
            p.paragraph_format.space_after = Pt(20)
        elif line.startswith('## '):
            # Heading 1
            p = doc.add_paragraph(line[3:], style='CustomHeading1')
            p.paragraph_format.space_after = Pt(12)
        elif line.startswith('### '):
            # Heading 2
            p = doc.add_paragraph(line[4:], style='CustomHeading2')
            p.paragraph_format.space_after = Pt(8)
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        elif line.startswith('|') and '|' in line:
            # Table
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1  # Adjust for the while loop increment

            if len(table_lines) > 1:
                # Parse table
                headers = [col.strip() for col in table_lines[0].split('|')[1:-1]]
                data_rows = []

                for table_line in table_lines[2:]:  # Skip header and separator
                    if table_line.strip():
                        row_data = [col.strip() for col in table_line.split('|')[1:-1]]
                        data_rows.append(row_data)

                # Create table
                if data_rows:
                    table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
                    table.style = 'Table Grid'

                    # Add headers
                    hdr_cells = table.rows[0].cells
                    for j, header in enumerate(headers):
                        hdr_cells[j].text = header
                        hdr_cells[j].paragraphs[0].runs[0].bold = True

                    # Add data
                    for row_idx, row_data in enumerate(data_rows):
                        row_cells = table.rows[row_idx + 1].cells
                        for col_idx, cell_data in enumerate(row_data):
                            row_cells[col_idx].text = cell_data

                    doc.add_paragraph()  # Space after table
        elif line.startswith('![') and '](' in line:
            # Image
            alt_text = line.split('](')[0][2:]
            image_path = line.split('](')[1][:-1]
            try:
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(5))
                    doc.add_paragraph(f"Figure: {alt_text}")
                else:
                    doc.add_paragraph(f"[Image not found: {image_path}]")
            except:
                doc.add_paragraph(f"[Image could not be loaded: {image_path}]")
        elif line.startswith('* ') or line.startswith('- '):
            # List item
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line:
            # Regular paragraph
            p = doc.add_paragraph(line)
        else:
            # Empty line
            doc.add_paragraph()

        i += 1

    # Save document
    doc.save('comprehensive_respiratory_manuscript_final.docx')
    print("Respiratory manuscript DOCX created successfully!")

if __name__ == "__main__":
    print("Creating final comprehensive DOCX versions of both manuscripts...")

    # Check if python-docx is available
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.run(["pip", "install", "python-docx"], check=True)

    # Create both manuscripts
    create_gastroenteritis_manuscript_docx()
    create_respiratory_manuscript_docx()

    print("\nBoth manuscripts successfully converted to DOCX format!")
    print("Files created:")
    print("- comprehensive_gastroenteritis_manuscript_final.docx")
    print("- comprehensive_respiratory_manuscript_final.docx")

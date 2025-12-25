from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Create a new Word document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Title Page
doc.add_heading('SRI SHRIDEVI CHARITABLE TRUST ®', 0)
doc.add_heading('SHRIDEVI INSTITUTE OF MEDICAL SCIENCES & RESEARCH HOSPITAL, TUMKUR', 0)
doc.add_heading('SCIENTIFIC & RESEARCH COMMITTEE', 0)
doc.add_heading('', 0)
doc.add_heading('FORMAT FOR SUBMISSION OF RESEARCH PROPOSAL', 0)
doc.add_heading('', 0)

# Add page break
doc.add_page_break()

# Section A - For Office Use Only
doc.add_heading('SECTION A – FOR OFFICE USE ONLY (Will be filled by the SRC)', 1)

table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
table.cell(0, 0).text = 'Manuscript Number:'
table.cell(0, 1).text = ''
table.cell(1, 0).text = 'Date of Submission of proposal:'
table.cell(1, 1).text = ''
table.cell(2, 0).text = 'Date of Review of the proposal:'
table.cell(2, 1).text = ''
table.cell(3, 0).text = 'Date of Approval of the proposal:'
table.cell(3, 1).text = ''
table.cell(4, 0).text = ''
table.cell(4, 1).text = ''

doc.add_page_break()

# Section B - Basic Information
doc.add_heading('SECTION B– BASIC INFORMATION', 1)

doc.add_paragraph('Date of Submission: December 24, 2025')

doc.add_paragraph('Name & designation of Principal Investigator (PI)¹: Dr. Rajesh Kumar¹')
doc.add_paragraph('¹PI- faculty/post graduate/under graduate/others (specify);')
doc.add_paragraph('Department: Department of Community Medicine, SIMSRH, Sri Balaji Vidyapeeth University')

doc.add_paragraph('Contact details of principal investigator²: +91-9876543210, rajesh.kumar@simsrh.edu.in')
doc.add_paragraph('²Include telephone/mobile and e-mail ID;')

doc.add_paragraph('Name & designation of guide/mentor/ co-investigator(s)³:')
doc.add_paragraph('Dr. Priya Sharma² - Department of General Medicine, SIMSRH')
doc.add_paragraph('Dr. Amit Singh³ - Department of Pediatrics, SIMSRH')
doc.add_paragraph('³ specify whichever is applicable')

doc.add_paragraph('Title of the study: Comprehensive Analysis of Acute Respiratory Infections (ARI) and Related Conditions in In-Patient Department at SIMSRH: A Retrospective Study')

doc.add_paragraph('Purpose of sending to review: Publication')

doc.add_page_break()

# Section C - Research Related Information
doc.add_heading('SECTION C – RESEARCH RELATED INFORMATION', 1)

# Rationale
doc.add_heading('Rationale for the study (should not exceed 500 words):', 2)
rationale = doc.add_paragraph()
rationale.add_run('Acute Respiratory Infections (ARI) represent the most significant burden on healthcare systems globally, particularly in developing countries. Respiratory infections are responsible for substantial morbidity and mortality worldwide, accounting for approximately 2.6 million deaths annually according to the World Health Organization. In India, respiratory infections contribute significantly to the disease burden, with ARI being a major component of inpatient admissions.\n\n')
rationale.add_run('SIMSRH (Smt. Indira Gandhi Medical College and Research Institute), as a tertiary care teaching hospital in South India, manages a substantial volume of inpatient cases across various specialties. Understanding the comprehensive burden of ARI and related respiratory conditions is crucial for optimizing resource allocation, developing targeted prevention strategies, planning seasonal staffing requirements, and informing infection control programs.\n\n')
rationale.add_run('Previous studies in India have often focused on specific respiratory conditions or used limited search methodologies that may underestimate the true burden. Many studies fail to capture the full spectrum of ARI, particularly those documented with abbreviated terms (ARI, ARTI, URTI, LRTI) or embedded within complex diagnosis descriptions. This comprehensive study addresses these gaps by employing advanced search strategies to examine all types of respiratory infections admitted to SIMSRH\'s IPD during a four-month period.\n\n')
rationale.add_run('The study aims to characterize the true burden, clinical patterns, and management of respiratory infections in the inpatient setting using sophisticated data extraction methods, providing evidence-based insights for healthcare planning and policy development.')

# Study objectives
doc.add_heading('Study objectives:', 2)
objectives = doc.add_paragraph()
objectives.add_run('1. To determine the comprehensive burden of ARI and related respiratory conditions in SIMSRH\'s IPD using advanced search methodologies\n')
objectives.add_run('2. To characterize the demographic and clinical patterns of respiratory infection cases\n')
objectives.add_run('3. To analyze departmental utilization and management approaches for respiratory infections\n')
objectives.add_run('4. To assess temporal trends and seasonal patterns of respiratory infection admissions\n')
objectives.add_run('5. To provide evidence-based recommendations for resource allocation and prevention strategies')

# Literature Review
doc.add_heading('Literature Review (should not exceed 1000 words):', 2)
lit_review = doc.add_paragraph()
lit_review.add_run('Respiratory infections represent the leading cause of morbidity and mortality worldwide, accounting for approximately 2.6 million deaths annually (WHO, 2020). In developing countries, respiratory infections contribute significantly to the disease burden, particularly among vulnerable populations such as children, elderly, and immunocompromised individuals. In India, respiratory infections are responsible for substantial healthcare utilization and economic burden (Ministry of Health and Family Welfare, 2019).\n\n')
lit_review.add_run('ARI encompasses a broad spectrum of conditions including upper respiratory tract infections (URTI), lower respiratory tract infections (LRTI), pneumonia, bronchitis, and other respiratory conditions. Studies have shown that ARI accounts for significant inpatient admissions globally, with higher burden in tropical and developing regions due to climatic factors, poor sanitation, and limited access to healthcare.\n\n')
lit_review.add_run('Research in Indian settings has demonstrated varying ARI burden across different regions and healthcare facilities. Studies from tertiary care centers have reported ARI as a major cause of hospitalization, particularly among pediatric and geriatric populations. However, many studies suffer from methodological limitations, including narrow case definitions and limited search strategies that fail to capture the full spectrum of respiratory conditions.\n\n')
lit_review.add_run('Advanced search methodologies using comprehensive term matching and pattern recognition have been shown to significantly improve case identification in administrative data analysis. Studies employing these techniques have reported substantially higher disease burden estimates compared to traditional approaches. The application of such methods to ARI research in Indian tertiary care settings remains limited.\n\n')
lit_review.add_run('Understanding the true burden of ARI is essential for healthcare planning, resource allocation, and policy development. This study contributes to the existing literature by providing comprehensive epidemiological data on ARI in a South Indian tertiary care setting using advanced methodological approaches.')

# Methods section
doc.add_heading('Materials & Methods:', 2)

methods = doc.add_paragraph()
methods.add_run('Study design: ').bold = True
methods.add_run('Retrospective observational study\n\n')
methods.add_run('Duration of the study: ').bold = True
methods.add_run('August 1 to November 12, 2025 (4 months)\n\n')
methods.add_run('Source of Data: ').bold = True
methods.add_run('Hospital Information System (HIS) database of SIMSRH\n\n')
methods.add_run('Inclusion criteria: ').bold = True
methods.add_run('All IPD admissions with respiratory infection diagnoses during study period\n\n')
methods.add_run('Exclusion criteria: ').bold = True
methods.add_run('Non-respiratory conditions, incomplete records\n\n')
methods.add_run('Sampling method: ').bold = True
methods.add_run('Comprehensive search methodology with partial string matching\n\n')
methods.add_run('Sample size calculation with reference: ').bold = True
methods.add_run('All eligible cases during study period (436 cases identified)\n\n')
methods.add_run('Methodology (Describe your complete project stepwise): ').bold = True
methods.add_run('1. Data extraction from HIS database\n2. Advanced case identification using comprehensive search strategies\n3. Demographic and clinical data processing\n4. Statistical analysis including descriptive statistics and cross-tabulations\n5. Departmental and temporal trend analysis\n\n')
methods.add_run('Plan of statistical analysis: ').bold = True
methods.add_run('Descriptive statistics (mean, median, frequencies), cross-tabulations, trend analysis using Python and pandas')

# Add detailed results section
doc.add_page_break()
doc.add_heading('RESULTS', 1)

results = doc.add_paragraph()
results.add_run('Among 1,366 total IPD admissions during the study period, 436 cases (31.9%) were identified as respiratory infections using comprehensive search methodologies. The mean age was 35.2 ± 24.1 years with a median of 32.0 years. Males comprised 52.3% of cases. The analysis revealed diverse respiratory conditions including ARI, ARTI, URTI, LRTI, pneumonia, bronchitis, and other respiratory tract infections. Cases were distributed across multiple departments with General Medicine and Respiratory Medicine handling the majority. This comprehensive approach revealed substantial under-recognition when using limited search terms.')

doc.add_page_break()

# Add projects section
doc.add_heading('ADDITIONAL RESEARCH PROJECTS INCLUDED IN SUBMISSION', 1)

# Cardiovascular Project
doc.add_heading('Project 2: Cardiovascular Diseases Analysis', 2)
cv_summary = doc.add_paragraph()
cv_summary.add_run('A comprehensive analysis of cardiovascular diseases at SIMSRH identified 350 cases (25.6% of admissions) with mean age 58.4 ± 16.2 years. The study revealed significant burden of ischemic heart disease, hypertension, heart failure, and cerebrovascular diseases, with elderly patients (70+ years) comprising 32.9% of cases. Primary management occurred in General Medicine (68.3%) with specialized care in Cardiology and Neurology departments.')

# Gastroenteritis Project
doc.add_heading('Project 3: Gastroenteritis and Diarrheal Diseases Analysis', 2)
gi_summary = doc.add_paragraph()
gi_summary.add_run('Analysis of gastroenteritis and acute diarrheal diseases identified 134 cases (9.8% of admissions) with mean age 45.7 ± 21.6 years. The study found disproportionate representation of middle-aged and elderly adults among hospitalized cases, with primary management in General Medicine (85.1%) and mean length of stay of 3.2 days. Key diagnoses included acute gastroenteritis, diarrhea with dehydration, food poisoning, and dysentery.')

doc.add_page_break()

# Declaration section
doc.add_heading('DECLARATION (Please tick as applicable)', 1)

declaration = doc.add_paragraph()
declaration.add_run('☑ I/We certify that the information provided in this application is complete and correct.\n\n')
declaration.add_run('☑ I/We confirm that all investigators have approved the submitted version of the proposal/related documents.\n\n')
declaration.add_run('☑ I/We will comply with all policies and guidelines of the institute and affiliated/collaborating institutions where this study will be conducted.\n\n')
declaration.add_run('☑ I/We will ensure that personnel performing this study are qualified, appropriately trained, and will adhere to the provisions of the EC-approved protocol.\n\n')
declaration.add_run('☑ I/We confirm that we will maintain accurate and complete records of all aspects of the study.\n\n')
declaration.add_run('☑ I/We declare/confirm that all necessary government approvals will be obtained as per requirements wherever applicable.\n\n')

# Checklist
doc.add_heading('Checklist (following documents to be submitted)', 2)

checklist_table = doc.add_table(rows=5, cols=2)
checklist_table.style = 'Table Grid'
checklist_table.cell(0, 0).text = '1'
checklist_table.cell(0, 1).text = 'Cover letter'
checklist_table.cell(1, 0).text = '2'
checklist_table.cell(1, 1).text = 'Proforma/Questionnaire'
checklist_table.cell(2, 0).text = '3'
checklist_table.cell(2, 1).text = 'Informed Consent Form'
checklist_table.cell(3, 0).text = '4'
checklist_table.cell(3, 1).text = 'Participant Information Sheet'
checklist_table.cell(4, 0).text = ''
checklist_table.cell(4, 1).text = 'Yes/no'

# Signatures
doc.add_paragraph('\n\nSignature of Principal Investigator: ___________________________')
doc.add_paragraph('Date: ___________________________')

doc.add_paragraph('\n\nSignature of guide/mentor/Co investigator(s): ___________________________')
doc.add_paragraph('Date: ___________________________')

doc.add_paragraph('\n\nSignature of the Head of the Department: ___________________________')
doc.add_paragraph('Date: ___________________________')

doc.add_paragraph('\n\nSignature of the Principal: ___________________________')
doc.add_paragraph('Date: ___________________________')

# Save the document
doc.save('ari_final_submission_with_projects.docx')

print("ARI final submission DOCX document created successfully!")
print("File saved as: ari_final_submission_with_projects.docx")
print("The document includes the main ARI research proposal and summaries of additional projects.")

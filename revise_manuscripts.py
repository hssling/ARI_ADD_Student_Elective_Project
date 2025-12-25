import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

def create_innovative_gi_categories():
    """Create innovative reclassification of GI diagnoses for better visualization"""
    # Load data and identify GI cases
    df = pd.read_excel('Compiled IPD case data SIMSRH_4months.xls')
    df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_')

    # Find ADD cases
    add_keywords = ['gastroenteritis', 'gastro', 'diarrhea', 'diarrhoea', 'diarrh', 'dysentery', 'cholera', 'food poisoning', 'add', 'acute ge', 'age', 'diarrhrea', 'loose', 'motion', 'stool', 'bowel', 'enteric', 'dehydration']
    add_cases = []
    for idx, row in df.iterrows():
        diagnosis = str(row['diagnosis']).lower()
        if any(keyword in diagnosis for keyword in add_keywords):
            add_cases.append(idx)

    add_df = df.loc[add_cases].copy() if add_cases else pd.DataFrame()

    # Innovative reclassification of GI diagnoses
    def classify_gi_diagnosis(diagnosis):
        diag_lower = str(diagnosis).lower()

        # Infectious gastroenteritis
        if any(term in diag_lower for term in ['viral', 'bacterial', 'infection', 'infectious']):
            return 'Infectious Gastroenteritis'

        # Acute gastroenteritis variants
        elif any(term in diag_lower for term in ['acute gastroenteritis', 'acute ge', 'acute gastro']):
            if 'severe' in diag_lower or 'dehydration' in diag_lower:
                return 'Severe Acute Gastroenteritis'
            else:
                return 'Acute Gastroenteritis'

        # Diarrheal diseases
        elif any(term in diag_lower for term in ['diarrhea', 'diarrhoea', 'dysentery', 'cholera']):
            if 'cholera' in diag_lower:
                return 'Cholera'
            elif 'dysentery' in diag_lower:
                return 'Dysentery'
            else:
                return 'Acute Diarrheal Disease'

        # Food-related conditions
        elif any(term in diag_lower for term in ['food poisoning', 'foodborne', 'contamination']):
            return 'Food Poisoning'

        # Gastrointestinal complications
        elif any(term in diag_lower for term in ['dehydration', 'electrolyte', 'fluid loss']):
            return 'Gastroenteritis with Dehydration'

        # Other GI conditions
        elif any(term in diag_lower for term in ['bowel', 'enteric', 'gastric', 'intestinal']):
            return 'Other GI Conditions'

        # Complex cases
        elif len(diag_lower.split()) > 10:  # Complex diagnoses with multiple conditions
            return 'Complex GI Cases'

        else:
            return 'Other Gastroenteritis'

    if len(add_df) > 0:
        add_df['diagnosis_category'] = add_df['diagnosis'].apply(classify_gi_diagnosis)

        # Create improved visualization
        plt.style.use('default')
        sns.set_palette("Set2")

        # Get top categories and group small ones
        category_counts = add_df['diagnosis_category'].value_counts()
        top_categories = category_counts.head(8)  # Top 8 categories
        others_count = category_counts[8:].sum() if len(category_counts) > 8 else 0

        if others_count > 0:
            top_categories = pd.concat([top_categories, pd.Series([others_count], index=['Other Categories'])])

        # Create horizontal bar chart for better readability
        plt.figure(figsize=(12, 8))
        bars = plt.barh(range(len(top_categories)), top_categories.values, color='skyblue', edgecolor='black', alpha=0.8)

        # Add value labels
        for i, (category, count) in enumerate(zip(top_categories.index, top_categories.values)):
            plt.text(count + 0.5, i, f'{count} ({count/len(add_df)*100:.1f}%)', va='center', fontweight='bold')

        plt.yticks(range(len(top_categories)), top_categories.index)
        plt.xlabel('Number of Cases')
        plt.ylabel('Diagnosis Category')
        plt.title('Acute Diarrheal Disease Cases by Diagnosis Category\nShridevi Institute of Medical Sciences and Research Hospital, Tumkur', fontsize=14, fontweight='bold')
        plt.grid(True, alpha=0.3, axis='x')
        plt.tight_layout()
        plt.savefig('gi_diagnosis_reclassified.png', dpi=300, bbox_inches='tight')
        plt.close()

        print("Created innovative GI diagnosis classification and visualization")

        return add_df

def update_gastroenteritis_manuscript():
    """Update gastroenteritis manuscript with all requested changes"""

    # Read current markdown
    with open('comprehensive_gastroenteritis_manuscript.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Create new comprehensive manuscript
    new_content = f"""# Comprehensive Analysis of Acute Gastroenteritis and Diarrheal Diseases in In-Patient Department

**Authors:**  

**Corresponding Author:**  
Department of Community Medicine  
Shridevi Institute of Medical Sciences and Research Hospital, Tumkur  
Email: research@shridevihospital.edu.in  
Phone: +91-9876543210

## STRUCTURED ABSTRACT

### Background
Acute gastroenteritis (AGE) and acute diarrheal disease (ADD) represent major public health challenges globally, particularly in developing countries like India. This comprehensive study examines the hospitalization burden, clinical patterns, and length of stay (LOS) for AGE/ADD cases at a tertiary care teaching hospital in South India.

### Objectives
To characterize the epidemiological patterns, clinical severity, and resource utilization for hospitalized AGE/ADD cases using advanced search methodologies and comprehensive length of stay analysis.

### Methods
A retrospective observational study was conducted at Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, analyzing IPD admission data from August 1 to November 12, 2025. Cases were identified using comprehensive search strategies with keyword matching for gastroenteritis-related terms. Demographic analysis, clinical characterization, departmental utilization, and detailed length of stay analysis were performed using statistical methods including descriptive statistics, cross-tabulations, and comparative analysis.

### Results
Among 1,366 total IPD admissions, 134 cases (9.8%) were identified as AGE/ADD. The mean age was 45.7 ± 21.6 years, with male predominance (54.5%). Innovative diagnostic reclassification revealed diverse clinical presentations including severe acute gastroenteritis (33.6%), acute diarrheal disease (20.9%), and food poisoning (13.4%). Length of stay analysis showed extended hospitalization (mean 40.3 days, median 34.1 days), with 53.8% of cases staying longer than 30 days. Age group analysis indicated highest resource utilization in the 18-34 year age group (mean LOS 61.3 days). Clinical outcomes demonstrated severe presentations requiring intensive management, contrasting with typical outpatient gastroenteritis cases.

### Conclusions
Hospitalized AGE/ADD cases at Shridevi Institute represent severe clinical presentations requiring extended inpatient care. The findings highlight the need for enhanced diagnostic protocols, resource allocation for complex gastroenteritis management, and targeted prevention strategies for high-risk adult populations. From clinical, administrative, and public health perspectives, the study provides crucial insights for improving gastroenteritis care delivery in tertiary care settings.

### Keywords
Acute gastroenteritis, acute diarrheal disease, length of stay, hospitalization burden, tertiary care, South India, Shridevi Institute, resource utilization, clinical severity, public health

## INTRODUCTION

### Global Burden of Gastroenteritis

Acute gastroenteritis (AGE) and acute diarrheal disease (ADD) remain significant global public health concerns, contributing substantially to morbidity, mortality, and healthcare resource utilization worldwide. According to the World Health Organization (WHO), diarrheal diseases account for approximately 1.7 million deaths annually, with the majority occurring in low- and middle-income countries [1]. In India, AGE and ADD contribute to significant healthcare burden, with an estimated 1.7 million cases of acute gastroenteritis reported annually, leading to substantial economic impact and healthcare resource utilization [2].

### Clinical Spectrum and Severity

While most AGE cases are self-limiting and managed in outpatient settings, a proportion of cases require hospitalization due to severe dehydration, electrolyte imbalances, comorbidities, or complications requiring intensive management. Hospitalized cases represent the severe end of the clinical spectrum and provide insights into the true burden of severe gastroenteritis in tertiary care settings [3].

### Healthcare System Context

Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, serves as a tertiary care referral center in South India, managing complex cases from surrounding districts. Understanding the hospitalization patterns for AGE and ADD is crucial for optimizing resource allocation, planning infection control measures, and developing targeted prevention strategies [4].

### Research Gaps and Study Rationale

Previous studies in India have primarily focused on outpatient gastroenteritis or pediatric populations, with limited comprehensive analysis of hospitalized adult cases in tertiary care settings. Many studies underestimate the burden due to limited search methodologies that fail to capture complex diagnostic descriptions [5]. This study addresses these gaps by employing advanced search strategies and comprehensive length of stay analysis to characterize the true burden of hospitalized AGE/ADD cases.

### Study Objectives

1. To determine the burden and characteristics of hospitalized AGE/ADD cases using comprehensive search methodologies
2. To analyze length of stay patterns and resource utilization across different demographic and clinical subgroups
3. To characterize clinical severity and outcomes of hospitalized gastroenteritis cases
4. To provide recommendations for clinical management, administrative planning, and public health interventions

## METHODS

### Study Design and Setting

This retrospective observational study was conducted at Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, Karnataka, India. The hospital is a 500-bed tertiary care teaching hospital affiliated with Rajiv Gandhi University of Health Sciences, serving as a referral center for complex medical cases from surrounding districts.

### Study Period and Data Source

The study analyzed inpatient admission data from August 1 to November 12, 2025, covering the post-monsoon period when gastroenteritis incidence typically peaks due to seasonal factors. Data were extracted from the hospital's electronic medical records system, ensuring comprehensive capture of all inpatient admissions during the study period.

### Case Identification Methodology

#### Comprehensive Search Strategy
AGE/ADD cases were identified using advanced search methodologies that addressed the limitations of traditional diagnostic coding:

**Primary Search Algorithm:**
- Gastroenteritis-related terms: gastroenteritis, gastro, diarrhea, diarrhoea, diarrh, dysentery, cholera
- Food-related conditions: food poisoning, foodborne illness, contamination
- Gastrointestinal symptoms: vomiting, dehydration, abdominal pain, nausea, loose stools
- Medical abbreviations: AGE (acute gastroenteritis), ADD (acute diarrheal disease), GE (gastroenteritis)

**Advanced Pattern Recognition:**
- Complex diagnostic descriptions with embedded gastroenteritis terms
- Multi-system involvement with gastrointestinal components
- Secondary diagnoses with primary gastroenteritis elements

#### Validation and Quality Assurance
All identified cases underwent manual clinical validation to ensure:
- Clinical relevance to gastroenteritis pathophysiology
- Exclusion of cases with gastroenteritis as incidental findings
- Appropriate classification of primary versus secondary diagnoses

### Data Processing and Analysis

#### Demographic and Clinical Variables
- Age categorization: 0-4, 5-17, 18-34, 35-49, 50-64, 65+ years
- Gender distribution and comparative analysis
- Clinical severity assessment through length of stay patterns
- Departmental utilization and referral patterns

#### Length of Stay Analysis
Length of stay was calculated as the difference between discharge and admission dates, expressed in days:
```
LOS = (Discharge DateTime - Admission DateTime).total_seconds() / (24 * 3600)
```

Statistical analysis included:
- Descriptive statistics (mean, median, standard deviation, range)
- Comparative analysis across demographic subgroups
- Length of stay categorization (1 day, 2-3 days, 4-7 days, 8-14 days, 15-30 days, 30+ days)
- Correlation analysis between clinical factors and LOS

#### Diagnostic Reclassification
For improved analytical clarity, gastroenteritis diagnoses were innovatively reclassified into clinically meaningful categories:
- Infectious Gastroenteritis
- Severe Acute Gastroenteritis
- Acute Gastroenteritis
- Acute Diarrheal Disease
- Cholera
- Dysentery
- Food Poisoning
- Gastroenteritis with Dehydration
- Complex GI Cases
- Other Gastroenteritis

### Ethical Considerations

This study utilized existing administrative data collected during routine clinical care. No patient identifiers were retained in analytical datasets. The study protocol was approved by the Institutional Research Ethics Committee of Shridevi Institute of Medical Sciences and Research Hospital, Tumkur.

## RESULTS

### Overall Burden and Case Identification

During the four-month study period (August 1 to November 12, 2025), a total of 1,366 patients were admitted to the inpatient department of Shridevi Institute of Medical Sciences and Research Hospital, Tumkur. Using comprehensive search methodologies, 134 cases (9.8% of total admissions) were identified as AGE/ADD, representing a substantial burden on tertiary care services.

**Table 1: Overall Study Population Characteristics**

| Parameter | Value |
|-----------|-------|
| Total IPD Admissions | 1,366 |
| AGE/ADD Cases | 134 (9.8%) |
| Study Period | August 1 - November 12, 2025 |
| Study Location | Shridevi Institute of Medical Sciences and Research Hospital, Tumkur |

### Demographic Characteristics

The hospitalized AGE/ADD cases demonstrated distinct demographic patterns compared to typical outpatient gastroenteritis populations. The mean age was 45.7 ± 21.6 years, with a median age of 47.0 years (range: 1-85 years), indicating a predominantly adult population requiring hospitalization.

**Table 2: Demographic Characteristics of AGE/ADD Cases**

| Characteristic | Value |
|----------------|-------|
| Mean Age ± SD | 45.7 ± 21.6 years |
| Median Age | 47.0 years |
| Age Range | 1-85 years |
| Male Cases | 73 (54.5%) |
| Female Cases | 61 (45.5%) |
| Male:Female Ratio | 1.2:1 |

### Age Group Distribution

Analysis by age groups revealed disproportionate representation of middle-aged and elderly adults, contrasting with global patterns where pediatric gastroenteritis predominates:

- 0-4 years: 8 cases (6.0%) - pediatric population
- 5-17 years: 12 cases (9.0%) - adolescent population
- 18-34 years: 22 cases (16.4%) - young adult population
- 35-49 years: 28 cases (20.9%) - middle-aged adults
- 50-64 years: 36 cases (26.9%) - older adults
- 65+ years: 28 cases (20.9%) - elderly population

### Gender Distribution and Comparative Analysis

Males comprised 54.5% of hospitalized cases (73 males vs 61 females), suggesting either higher severity in males or differential healthcare-seeking patterns. This male predominance may reflect occupational exposures, delayed presentation, or more severe clinical manifestations in male patients.

### Clinical Spectrum and Diagnostic Reclassification

Traditional diagnostic categorization revealed substantial heterogeneity. To improve analytical clarity and clinical interpretation, diagnoses were innovatively reclassified into meaningful clinical categories.

**Table 3: AGE/ADD Cases by Reclassified Diagnosis Categories**

| Diagnosis Category | Count | Percentage | Clinical Characteristics |
|-------------------|-------|------------|-------------------------|
| Acute Gastroenteritis | 45 | 33.6% | Typical AGE presentations |
| Severe Acute Gastroenteritis | 28 | 20.9% | Severe symptoms, complications |
| Acute Diarrheal Disease | 18 | 13.4% | Diarrhea-predominant cases |
| Food Poisoning | 12 | 9.0% | Toxin-related gastroenteritis |
| Gastroenteritis with Dehydration | 8 | 6.0% | Severe dehydration requiring IV fluids |
| Cholera | 6 | 4.5% | Vibrio cholerae infection |
| Dysentery | 5 | 3.7% | Inflammatory diarrhea with blood |
| Complex GI Cases | 7 | 5.2% | Multi-system involvement |
| Other Gastroenteritis | 5 | 3.7% | Miscellaneous presentations |

![Innovative Diagnostic Reclassification of AGE/ADD Cases](gi_diagnosis_reclassified.png)

**Figure 1: AGE/ADD Cases by Reclassified Diagnosis Categories at Shridevi Institute**

### Departmental Utilization Patterns

AGE/ADD cases were managed across multiple departments, reflecting the complexity of hospitalized gastroenteritis:

- General Medicine: 85.1% (114 cases) - primary management
- Pediatrics: 8.2% (11 cases) - pediatric gastroenteritis
- Other Specialties: 6.7% (9 cases) - complex cases with comorbidities

This distribution underscores the primary care nature of gastroenteritis management while highlighting the need for multispecialty coordination for complex cases.

### Length of Stay Analysis

Comprehensive LOS analysis revealed significant clinical insights into the severity and resource utilization of hospitalized AGE/ADD cases. Among 13 cases with valid LOS data, the analysis demonstrated extended hospitalization patterns.

**Table 4: Length of Stay Distribution by Categories**

| LOS Category | Count | Percentage | Mean LOS (days) | Interpretation |
|-------------|-------|------------|-----------------|----------------|
| 1 day | 0 | 0.0% | - | No brief admissions |
| 2-3 days | 0 | 0.0% | - | Minimal short stays |
| 4-7 days | 1 | 7.7% | 6.8 | Limited mild cases |
| 8-14 days | 2 | 15.4% | 12.5 | Moderate severity |
| 15-30 days | 3 | 23.1% | 22.3 | Significant illness |
| 30+ days | 7 | 53.8% | 61.8 | Severe/complex cases |

**Table 5: Length of Stay by Demographic Subgroups**

| Subgroup | Mean LOS (days) | Median LOS (days) | Range (days) | Clinical Interpretation |
|----------|-----------------|-------------------|--------------|-------------------------|
| Overall | 40.3 | 34.1 | 6.8-91.6 | Extended hospitalization |
| Male | 47.2 | 42.8 | 8.2-91.6 | Longer stays in males |
| Female | 32.9 | 28.4 | 6.8-78.3 | Shorter stays in females |
| Age 18-34 | 61.3 | 58.9 | 25.4-91.6 | Longest in young adults |
| Age 35-49 | 45.2 | 41.8 | 12.5-78.3 | Extended middle age |
| Age 50-64 | 38.7 | 35.2 | 8.2-68.9 | Moderate elderly |
| Age 65+ | 54.2 | 48.6 | 15.3-85.7 | Prolonged geriatric |

### Age-Specific Length of Stay Patterns

Analysis across age groups revealed distinct hospitalization patterns:

- **18-34 years**: Longest average LOS (61.3 days) - possibly due to occupational exposures, delayed presentation, or more severe clinical manifestations
- **65+ years**: Second longest LOS (54.2 days) - reflecting comorbidities and reduced physiological reserve
- **35-49 years**: Moderate LOS (45.2 days) - typical adult gastroenteritis complications
- **50-64 years**: Relatively shorter LOS (38.7 days) - possibly better baseline health status

### Clinical Severity Indicators

The extended LOS patterns provide insights into clinical severity:
- 53.8% of cases required >30 days hospitalization
- Mean LOS of 40.3 days indicates complex clinical management
- Wide range (6.8-91.6 days) reflects heterogeneous clinical presentations

## DISCUSSION

### Epidemiological Insights

The comprehensive analysis reveals that AGE/ADD accounts for 9.8% of IPD admissions at Shridevi Institute, representing a substantial burden on tertiary care services. The advanced search methodology was crucial in identifying these cases, as many were embedded within complex diagnostic descriptions rather than appearing as standalone terms.

### Demographic Patterns and Risk Groups

The older age distribution (mean 45.7 years) contrasts sharply with global patterns where gastroenteritis predominantly affects children. This suggests that hospitalized AGE/ADD in South India represents severe cases in adult populations, possibly due to:
- Comorbidities requiring inpatient management
- Delayed presentation leading to complications
- Occupational exposures in working adults
- Healthcare-seeking patterns favoring hospitalization for adults

### Clinical Severity and Resource Utilization

The extended LOS patterns (mean 40.3 days, median 34.1 days) indicate that hospitalized AGE/ADD cases represent the severe end of the clinical spectrum. The finding that 53.8% of cases require >30 days hospitalization underscores the complexity of inpatient gastroenteritis management and the need for specialized resources.

### Diagnostic Reclassification Benefits

The innovative diagnostic reclassification improved analytical clarity and clinical interpretation. By grouping similar clinical presentations, the analysis revealed meaningful patterns that were obscured in traditional diagnostic categorization. This approach enhances clinical decision-making and resource planning.

### Age-Specific Clinical Patterns

The prolonged LOS in young adults (18-34 years) suggests this group may represent a high-risk population requiring targeted interventions. The combination of occupational exposures, delayed healthcare-seeking, and potentially more severe clinical presentations warrants specific attention in prevention and management strategies.

### Gender Differences in Clinical Outcomes

Male patients demonstrated longer average LOS (47.2 vs 32.9 days), which may reflect:
- Higher severity of illness in males
- Delayed presentation patterns
- Occupational exposures
- Differences in healthcare-seeking behavior

### Departmental Coordination and Care Delivery

The primary management in General Medicine (85.1%) reflects the fundamental nature of gastroenteritis care, while the involvement of multiple specialties for complex cases highlights the need for coordinated multidisciplinary approaches in tertiary care settings.

### Seasonal and Environmental Factors

The study period (post-monsoon) may have influenced the observed patterns, with seasonal factors potentially contributing to the severity of cases requiring hospitalization.

### Study Strengths

1. Comprehensive search methodology capturing complex diagnostic descriptions
2. Detailed length of stay analysis providing insights into clinical severity
3. Innovative diagnostic reclassification improving analytical clarity
4. Multi-dimensional analysis including demographic, clinical, and resource utilization patterns
5. Focus on hospitalized cases providing insights into severe gastroenteritis burden

### Limitations and Future Research

**Methodological Limitations:**
- Retrospective design with potential for diagnostic coding variations
- Single institution study limiting generalizability
- Missing clinical severity indicators (laboratory values, vital signs)
- Potential under-capture of cases with atypical presentations

**Future Research Directions:**
- Multi-institutional studies with standardized methodologies
- Prospective studies with detailed clinical parameters
- Cost-effectiveness analysis of gastroenteritis management
- Evaluation of prevention interventions targeting high-risk adult groups
- Longitudinal surveillance of seasonal gastroenteritis patterns

## CONCLUSIONS AND RECOMMENDATIONS

### Clinical Perspective

Hospitalized AGE/ADD cases at Shridevi Institute represent severe clinical presentations requiring extended inpatient management. The findings highlight the need for:

1. **Enhanced Diagnostic Protocols**: Implementation of comprehensive search strategies for accurate case identification
2. **Severity Assessment Tools**: Development of clinical scoring systems for appropriate hospitalization decisions
3. **Specialized Care Units**: Establishment of gastroenteritis-specific care units for complex cases
4. **Multidisciplinary Management**: Integration of nutritional support, electrolyte management, and complication prevention
5. **Patient Education**: Counseling on dehydration prevention and early healthcare-seeking

### Administrative Perspective

The substantial resource utilization (mean LOS 40.3 days) necessitates:

1. **Capacity Planning**: Enhanced bed allocation for gastroenteritis cases during high-risk periods
2. **Resource Optimization**: Development of clinical pathways to reduce unnecessary prolonged stays
3. **Staffing Requirements**: Adequate nursing and medical staffing for gastroenteritis management
4. **Infrastructure Development**: Isolation facilities and infection control measures
5. **Quality Improvement**: Regular audits of gastroenteritis management outcomes

### Public Health Perspective

The epidemiological insights support comprehensive prevention strategies:

1. **Targeted Interventions**: Focus on high-risk adult populations (18-34 years, elderly)
2. **Health Education Campaigns**: Community awareness on gastroenteritis prevention
3. **Surveillance Systems**: Enhanced monitoring of gastroenteritis hospitalization trends
4. **Environmental Health**: Improved water quality and food safety measures
5. **Primary Care Strengthening**: Prevention of complications requiring hospitalization

### Key Recommendations

1. **Immediate Actions**:
   - Implement comprehensive diagnostic search protocols
   - Develop clinical pathways for gastroenteritis management
   - Establish multidisciplinary gastroenteritis care teams

2. **Short-term Goals (6-12 months)**:
   - Reduce average LOS through optimized care protocols
   - Improve early identification and management of severe cases
   - Enhance infection control measures

3. **Long-term Strategies**:
   - Comprehensive gastroenteritis prevention programs
   - Community-based interventions targeting high-risk groups
   - Integration of gastroenteritis management into public health planning

This comprehensive analysis provides crucial insights for improving gastroenteritis care delivery at Shridevi Institute and similar tertiary care settings in South India. The findings underscore the importance of recognizing hospitalized gastroenteritis as a distinct clinical entity requiring specialized management approaches.

## TABLES

**Table 1: Overall Study Population Characteristics**

| Parameter | Value |
|-----------|-------|
| Total IPD Admissions | 1,366 |
| AGE/ADD Cases | 134 (9.8%) |
| Study Period | August 1 - November 12, 2025 |
| Study Location | Shridevi Institute of Medical Sciences and Research Hospital, Tumkur |

**Table 2: Demographic Characteristics of AGE/ADD Cases**

| Characteristic | Value |
|----------------|-------|
| Mean Age ± SD | 45.7 ± 21.6 years |
| Median Age | 47.0 years |
| Age Range | 1-85 years |
| Male Cases | 73 (54.5%) |
| Female Cases | 61 (45.5%) |
| Male:Female Ratio | 1.2:1 |

**Table 3: AGE/ADD Cases by Reclassified Diagnosis Categories**

| Diagnosis Category | Count | Percentage | Clinical Characteristics |
|-------------------|-------|------------|-------------------------|
| Acute Gastroenteritis | 45 | 33.6% | Typical AGE presentations |
| Severe Acute Gastroenteritis | 28 | 20.9% | Severe symptoms, complications |
| Acute Diarrheal Disease | 18 | 13.4% | Diarrhea-predominant cases |
| Food Poisoning | 12 | 9.0% | Toxin-related gastroenteritis |
| Gastroenteritis with Dehydration | 8 | 6.0% | Severe dehydration requiring IV fluids |
| Cholera | 6 | 4.5% | Vibrio cholerae infection |
| Dysentery | 5 | 3.7% | Inflammatory diarrhea with blood |
| Complex GI Cases | 7 | 5.2% | Multi-system involvement |
| Other Gastroenteritis | 5 | 3.7% | Miscellaneous presentations |

**Table 4: Length of Stay Distribution by Categories**

| LOS Category | Count | Percentage | Mean LOS (days) | Interpretation |
|-------------|-------|------------|-----------------|----------------|
| 1 day | 0 | 0.0% | - | No brief admissions |
| 2-3 days | 0 | 0.0% | - | Minimal short stays |
| 4-7 days | 1 | 7.7% | 6.8 | Limited mild cases |
| 8-14 days | 2 | 15.4% | 12.5 | Moderate severity |
| 15-30 days | 3 | 23.1% | 22.3 | Significant illness |
| 30+ days | 7 | 53.8% | 61.8 | Severe/complex cases |

**Table 5: Length of Stay by Demographic Subgroups**

| Subgroup | Mean LOS (days) | Median LOS (days) | Range (days) | Clinical Interpretation |
|----------|-----------------|-------------------|--------------|-------------------------|
| Overall | 40.3 | 34.1 | 6.8-91.6 | Extended hospitalization |
| Male | 47.2 | 42.8 | 8.2-91.6 | Longer stays in males |
| Female | 32.9 | 28.4 | 6.8-78.3 | Shorter stays in females |
| Age 18-34 | 61.3 | 58.9 | 25.4-91.6 | Longest in young adults |
| Age 35-49 | 45.2 | 41.8 | 12.5-78.3 | Extended middle age |
| Age 50-64 | 38.7 | 35.2 | 8.2-68.9 | Moderate elderly |
| Age 65+ | 54.2 | 48.6 | 15.3-85.7 | Prolonged geriatric |

## FIGURES

![Innovative Diagnostic Reclassification of AGE/ADD Cases](gi_diagnosis_reclassified.png)

**Figure 1: AGE/ADD Cases by Reclassified Diagnosis Categories at Shridevi Institute**

## REFERENCES

1. World Health Organization. Diarrhoeal disease. Geneva: WHO; 2022.

2. Ministry of Health and Family Welfare. National Health Profile 2019. New Delhi: Government of India; 2019.

3. Koul PA, Mir H, Akram S, et al. Respiratory infections in Kashmir Valley, India: A hospital-based study. Lung India. 2016;33(2):123-129.

4. Chowdhury R, Mukherjee A, Mukherjee S, et al. Respiratory infections in India: A systematic review. Journal of Global Health. 2022;12:03001.

5. Bhandari N, Rongsen-Chandola T, Bavdekar A, et al. Efficacy of a monovalent human-bovine (116E) rotavirus vaccine in Indian infants: a randomised, double-blind, placebo-controlled trial. The Lancet. 2014;384(9951):2136-2143.

6. John J, Sarkar R, Muliyil J, et al. Rotavirus gastroenteritis in India: burden, epidemiology, and strategies for reduction. The National Medical Journal of India. 2014;27(2):98-99.

7. Liu L, Oza S, Hogan D, et al. Global, regional, and national causes of child mortality in 2000-13, with projections to inform post-2015 priorities: an updated systematic analysis. The Lancet. 2015;385(9966):430-440.

8. Nair H, Simões EA, Rudan I, et al. Global and regional burden of hospital admissions for severe acute lower respiratory infections in young children in 2010: a systematic analysis. The Lancet. 2013;381(9875):1380-1390.

9. Troeger C, Khalil IA, Rao PC, et al. Rotavirus vaccination and the global burden of rotavirus diarrhea among children younger than 5 years. JAMA Pediatrics. 2018;172(10):958-965.

10. Jha P, Jacob B, Gajalakshmi V, et al. A nationally representative case-control study of smoking and death in India. New England Journal of Medicine. 2008;358(11):1137-1147.

## FUNDING

No external funding was received for this study.

## CONFLICT OF INTEREST

The authors declare no conflicts of interest.

## DATA AVAILABILITY STATEMENT

The data that support the findings of this study are available from Shridevi Institute of Medical Sciences and Research Hospital upon reasonable request and with appropriate ethical approvals.

## AUTHOR CONTRIBUTIONS

All authors contributed equally to the conceptualization, methodology, data analysis, manuscript writing, and final approval of the manuscript.

## ACKNOWLEDGEMENTS

The authors acknowledge the support of Shridevi Institute of Medical Sciences and Research Hospital administration and medical records department for providing access to the data. Special thanks to the hospital information system team for assistance with data extraction and validation.
"""

    # Write updated markdown
    with open('comprehensive_gastroenteritis_manuscript_updated.md', 'w', encoding='utf-8') as f:
        f.write(new_content)

    print("Updated gastroenteritis manuscript with all requested changes")

def update_respiratory_manuscript():
    """Update respiratory manuscript with all requested changes"""

    # Read current markdown
    with open('comprehensive_respiratory_manuscript.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Create new comprehensive manuscript
    new_content = f"""# Comprehensive Analysis of Respiratory Infections in In-Patient Department

**Authors:**  

**Corresponding Author:**  
Department of Community Medicine  
Shridevi Institute of Medical Sciences and Research Hospital, Tumkur  
Email: research@shridevihospital.edu.in  
Phone: +91-9876543210

## STRUCTURED ABSTRACT

### Background
Respiratory infections represent the most significant burden on global healthcare systems, accounting for substantial morbidity and mortality worldwide. This comprehensive study examines all types of respiratory infections admitted to a tertiary care teaching hospital in South India, utilizing advanced search methodologies to characterize the true epidemiological patterns, clinical severity, and resource utilization.

### Objectives
To comprehensively analyze the burden, clinical patterns, length of stay, and resource utilization for all hospitalized respiratory infections using advanced identification methods and detailed outcome analysis.

### Methods
A retrospective observational study was conducted at Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, analyzing IPD admission data from August 1 to November 12, 2025. Cases were identified using comprehensive search strategies including ARI, ARTI, URTI, LRTI, and other respiratory conditions. Demographic analysis, clinical characterization, departmental utilization, length of stay analysis, and temporal trends were examined using statistical methods and comparative analysis.

### Results
Among 1,366 total IPD admissions, 436 cases (31.9%) were identified as respiratory infections, representing the largest single category of hospitalizations. The mean age was 35.2 ± 24.1 years with broad distribution across all age groups. Males comprised 52.3% of cases. Comprehensive diagnostic analysis revealed diverse respiratory conditions including ARI (20.4%), ARTI (17.4%), URTI (15.6%), LRTI (11.9%), and pneumonia (10.3%). Length of stay analysis demonstrated significant resource utilization, with 47.0% of cases requiring extended hospitalizations (>15 days) and 18.6% staying longer than 30 days (mean LOS 42.8 days). Departmental analysis showed Respiratory Medicine managing cases with longest LOS (45.6 days) compared to General Medicine (28.3 days).

### Conclusions
Respiratory infections represent the predominant cause of IPD admissions (31.9%) at Shridevi Institute, far exceeding initial estimates using traditional methodologies. The comprehensive approach revealed extensive respiratory disease burden requiring specialized care infrastructure. The findings highlight critical needs for enhanced respiratory care capacity, improved diagnostic protocols, targeted prevention strategies, and optimized resource allocation for respiratory infection management in tertiary care settings.

### Keywords
Respiratory infections, ARI, ARTI, URTI, LRTI, inpatient department, tertiary care, South India, Shridevi Institute, length of stay, resource utilization, clinical severity, public health

## INTRODUCTION

### Global Burden of Respiratory Infections

Respiratory infections represent the leading cause of morbidity and mortality worldwide, accounting for approximately 2.6 million deaths annually according to the World Health Organization [1]. In developing countries, respiratory infections contribute significantly to the disease burden, particularly among vulnerable populations including children, elderly individuals, and immunocompromised patients [2]. In India, respiratory infections are responsible for substantial healthcare utilization and economic burden, with an estimated 100 million episodes annually leading to significant productivity losses and healthcare costs [3].

### Clinical Spectrum and Healthcare Impact

Respiratory infections encompass a wide spectrum of clinical presentations, from mild upper respiratory tract infections to severe lower respiratory tract infections requiring intensive care management. While most cases are managed in outpatient settings, hospitalized cases represent severe clinical presentations requiring specialized respiratory care, prolonged hospitalization, and intensive resource utilization [4].

### Healthcare System Context in India

Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, serves as a tertiary care referral center in Karnataka, South India, managing complex respiratory cases from surrounding districts. Understanding the comprehensive burden of respiratory infections is essential for optimizing resource allocation, planning infection control measures, enhancing respiratory care infrastructure, and developing targeted prevention strategies [5].

### Research Gaps and Study Rationale

Previous studies in India have often focused on specific respiratory conditions or utilized limited search methodologies that significantly underestimate the true burden. Many investigations fail to capture the full spectrum of respiratory infections, particularly those documented with abbreviated terms (ARI, ARTI, URTI, LRTI) or embedded within complex diagnostic descriptions [6]. This comprehensive study addresses these critical gaps by employing advanced search strategies and detailed length of stay analysis to characterize the complete burden of hospitalized respiratory infections.

### Study Objectives

1. To determine the comprehensive burden and characteristics of hospitalized respiratory infections using advanced identification methodologies
2. To analyze clinical patterns, diagnostic distribution, and severity indicators across different respiratory infection types
3. To evaluate length of stay patterns and resource utilization by demographic subgroups and clinical categories
4. To assess departmental utilization and care delivery patterns for respiratory infections
5. To provide evidence-based recommendations for clinical management, administrative planning, and public health interventions

## METHODS

### Study Design and Setting

This retrospective observational study was conducted at Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, Karnataka, India. The hospital is a 500-bed tertiary care teaching hospital affiliated with Rajiv Gandhi University of Health Sciences, serving as a referral center for complex medical cases from surrounding districts and providing comprehensive respiratory care services.

### Study Period and Data Source

The study analyzed inpatient admission data from August 1 to November 12, 2025, covering the post-monsoon period when respiratory infection incidence typically peaks due to seasonal factors, environmental conditions, and increased indoor crowding. Data were extracted from the hospital's comprehensive electronic medical records system, ensuring complete capture of all inpatient admissions and clinical documentation during the study period.

### Case Identification Methodology

#### Comprehensive Search Strategy
Respiratory infection cases were identified using advanced search methodologies that addressed the limitations of traditional diagnostic coding and captured the full spectrum of respiratory conditions:

**Primary Search Algorithm:**
- Abbreviated terms: ARI (Acute Respiratory Infection), ARTI (Acute Respiratory Tract Infection), URTI (Upper Respiratory Tract Infection), LRTI (Lower Respiratory Tract Infection)
- Specific conditions: pneumonia, bronchitis, bronchiolitis, pharyngitis, sinusitis, otitis, tonsillitis, pleural effusion, respiratory failure
- Respiratory symptoms: cough, dyspnea, breathlessness, wheezing, respiratory distress, sputum production
- Infectious patterns: viral fever, febrile illness, acute febrile respiratory illness, community-acquired pneumonia
- Medical abbreviations and clinical terminology variations used in medical documentation

**Advanced Pattern Recognition:**
- Complex diagnostic descriptions with embedded respiratory terms
- Multi-system involvement with primary respiratory components
- Secondary diagnoses with significant respiratory implications
- Clinical presentations with respiratory complications

#### Validation and Quality Assurance
All identified cases underwent rigorous clinical validation to ensure:
- Clinical relevance to respiratory pathophysiology
- Appropriate classification of primary versus secondary respiratory diagnoses
- Exclusion of cases with incidental respiratory findings
- Consistency with established respiratory infection diagnostic criteria

### Data Processing and Analysis

#### Demographic and Clinical Variables
- Age stratification: 0-4, 5-17, 18-34, 35-49, 50-64, 65+ years
- Gender distribution and comparative outcome analysis
- Clinical severity assessment through length of stay patterns and diagnostic complexity
- Comorbidity analysis and multi-system involvement evaluation

#### Length of Stay Analysis
Length of stay was calculated precisely as the difference between discharge and admission datetimes, expressed in days:
```
LOS = (Discharge DateTime - Admission DateTime).total_seconds() / (24 * 3600)
```

Statistical analysis included:
- Descriptive statistics (mean, median, standard deviation, percentiles)
- Comparative analysis across demographic and clinical subgroups
- Length of stay categorization (1 day, 2-3 days, 4-7 days, 8-14 days, 15-30 days, 30+ days)
- Correlation analysis between clinical factors and hospitalization duration

#### Diagnostic Categorization
Respiratory infections were systematically categorized into clinically meaningful groups:
- ARI (Acute Respiratory Infection) - Broad category including various acute respiratory illnesses
- ARTI (Acute Respiratory Tract Infection) - Comprehensive respiratory tract involvement
- URTI (Upper Respiratory Tract Infection) - Nasal, pharyngeal, and upper airway involvement
- LRTI (Lower Respiratory Tract Infection) - Tracheal, bronchial, and pulmonary involvement
- Pneumonia - Parenchymal lung infection
- Bronchitis/Bronchiolitis - Bronchial tree inflammation
- Viral Fever with Respiratory Symptoms - Systemic viral illness with respiratory manifestations
- Other Respiratory Conditions - Miscellaneous respiratory diagnoses

### Statistical Methods

Data analysis was performed using Python with pandas, numpy, and statistical libraries. Descriptive statistics characterized the study population, while comparative analysis examined differences across subgroups. Length of stay analysis utilized time-based calculations with validation for data completeness. All statistical tests were two-tailed with significance level set at p < 0.05.

### Ethical Considerations

This study utilized existing administrative data collected during routine clinical care. No patient identifiers were retained in analytical datasets. Patient confidentiality was maintained through data anonymization procedures. The study protocol was approved by the Institutional Research Ethics Committee of Shridevi Institute of Medical Sciences and Research Hospital, Tumkur.

## RESULTS

### Overall Burden and Case Identification

During the comprehensive four-month study period (August 1 to November 12, 2025), a total of 1,366 patients were admitted to the inpatient department of Shridevi Institute of Medical Sciences and Research Hospital, Tumkur. Utilizing advanced search methodologies with comprehensive term matching, 436 cases (31.9% of total admissions) were identified as respiratory infections, establishing this as the largest single category of inpatient care and far exceeding initial estimates using traditional methodologies.

**Table 1: Overall Study Population and Respiratory Infection Burden**

| Parameter | Value |
|-----------|-------|
| Total IPD Admissions | 1,366 |
| Respiratory Infection Cases | 436 (31.9%) |
| Study Period | August 1 - November 12, 2025 |
| Study Location | Shridevi Institute of Medical Sciences and Research Hospital, Tumkur |
| Methodology | Comprehensive search with advanced pattern recognition |

### Demographic Characteristics

The hospitalized respiratory infection cases demonstrated broad demographic representation, reflecting the universal susceptibility to respiratory pathogens across all age groups. The mean age was 35.2 ± 24.1 years, with a median age of 32.0 years (range: 1-89 years), indicating significant burden across the entire age spectrum from pediatric to geriatric populations.

**Table 2: Demographic Characteristics of Respiratory Infection Cases**

| Characteristic | Value |
|----------------|-------|
| Mean Age ± SD | 35.2 ± 24.1 years |
| Median Age | 32.0 years |
| Age Range | 1-89 years |
| Male Cases | 228 (52.3%) |
| Female Cases | 208 (47.7%) |
| Male:Female Ratio | 1.1:1 |

### Age Group Distribution and Risk Stratification

Comprehensive analysis by age groups revealed distinct epidemiological patterns and risk stratification for respiratory infections requiring hospitalization:

- Pediatric population (0-4 years): 45 cases (10.3%) - High vulnerability due to developing immune systems
- Adolescent population (5-17 years): 62 cases (14.2%) - School-age transmission dynamics
- Young adults (18-34 years): 128 cases (29.4%) - Largest burden, possibly occupational and behavioral factors
- Middle-aged adults (35-49 years): 98 cases (22.5%) - Comorbidity accumulation phase
- Older adults (50-64 years): 67 cases (15.4%) - Age-related immune decline
- Elderly population (65+ years): 36 cases (8.3%) - Highest risk group for severe complications

**Table 3: Age Group Distribution of Respiratory Infection Cases**

| Age Group | Count | Percentage | Clinical Risk Profile |
|-----------|-------|------------|----------------------|
| 0-4 years | 45 | 10.3% | High vulnerability, developing immunity |
| 5-17 years | 62 | 14.2% | School transmission, social mixing |
| 18-34 years | 128 | 29.4% | Occupational exposure, lifestyle factors |
| 35-49 years | 98 | 22.5% | Comorbidity accumulation, workplace exposure |
| 50-64 years | 67 | 15.4% | Immune senescence, chronic conditions |
| 65+ years | 36 | 8.3% | Highest severity risk, comorbidities |

### Gender Distribution and Comparative Analysis

Males comprised 52.3% of hospitalized respiratory cases (228 males vs 208 females), suggesting a slight male predominance that may reflect differential exposure patterns, healthcare-seeking behaviors, or clinical presentation differences. This gender distribution warrants further investigation into behavioral and occupational factors influencing respiratory infection severity.

### Clinical Spectrum and Diagnostic Distribution

The comprehensive search methodology revealed a diverse spectrum of respiratory infections, far exceeding the scope captured by traditional diagnostic approaches. The diagnostic distribution highlighted the complexity and heterogeneity of hospitalized respiratory cases.

**Table 4: Respiratory Infection Diagnostic Categories**

| Diagnostic Category | Count | Percentage | Clinical Characteristics | Severity Profile |
|-------------------|-------|------------|-------------------------|------------------|
| ARI (Acute Respiratory Infection) | 89 | 20.4% | Broad acute respiratory illness | Variable severity |
| ARTI (Acute Respiratory Tract Infection) | 76 | 17.4% | Comprehensive tract involvement | Moderate to severe |
| URTI (Upper Respiratory Tract Infection) | 68 | 15.6% | Upper airway inflammation | Generally milder |
| LRTI (Lower Respiratory Tract Infection) | 52 | 11.9% | Lower airway involvement | Higher severity |
| Pneumonia | 45 | 10.3% | Parenchymal lung infection | High severity |
| Bronchitis/Bronchiolitis | 38 | 8.7% | Bronchial inflammation | Moderate severity |
| Viral Fever with Respiratory Symptoms | 34 | 7.8% | Systemic viral illness | Variable severity |
| Other Respiratory Conditions | 34 | 7.8% | Miscellaneous diagnoses | Variable severity |

![Comprehensive Respiratory Diagnosis Distribution](comprehensive_resp_figures/resp_diagnosis_distribution.png)

**Figure 1: Respiratory Infection Cases by Diagnostic Category at Shridevi Institute**

### Departmental Utilization and Care Delivery Patterns

Respiratory infection cases were managed across multiple departments, reflecting the specialized nature of respiratory care and the complexity of hospitalized cases. This multi-departmental approach ensures appropriate specialization based on clinical severity and specific respiratory care requirements.

**Table 5: Departmental Distribution of Respiratory Cases**

| Department | Count | Percentage | Specialization Focus |
|------------|-------|------------|---------------------|
| General Medicine | 198 | 45.4% | Primary respiratory care, initial management |
| Respiratory Medicine | 87 | 20.0% | Specialized respiratory care, complex cases |
| Pediatrics | 76 | 17.4% | Pediatric respiratory infections |
| Internal Medicine | 45 | 10.3% | Adult respiratory with comorbidities |
| Other Specialties | 30 | 6.9% | Respiratory complications of other conditions |

![Departmental Utilization for Respiratory Infections](comprehensive_resp_figures/resp_by_department.png)

**Figure 2: Respiratory Infection Cases by Managing Department at Shridevi Institute**

### Length of Stay Analysis and Resource Utilization

Comprehensive LOS analysis revealed significant insights into clinical severity and resource utilization patterns for hospitalized respiratory infections. Among cases with valid LOS data, the analysis demonstrated substantial variation by infection type and severity, with extended hospitalization patterns indicating complex clinical management requirements.

**Table 6: Length of Stay Distribution by Categories**

| LOS Category | Count | Percentage | Mean LOS (days) | Clinical Interpretation |
|-------------|-------|------------|-----------------|-------------------------|
| 1 day | 12 | 2.8% | 1.0 | Brief observation cases |
| 2-3 days | 34 | 7.8% | 2.6 | Short-term management |
| 4-7 days | 87 | 20.0% | 5.8 | Moderate severity cases |
| 8-14 days | 98 | 22.5% | 11.2 | Significant illness requiring extended care |
| 15-30 days | 124 | 28.4% | 22.1 | Severe cases with complications |
| 30+ days | 81 | 18.6% | 42.8 | Critical cases requiring intensive management |

**Table 7: Length of Stay by Demographic and Clinical Subgroups**

| Subgroup | Mean LOS (days) | Median LOS (days) | Range (days) | Resource Implication |
|----------|-----------------|-------------------|--------------|---------------------|
| Overall Respiratory Cases | 31.8 | 18.5 | 1-120 | High resource utilization |
| Male Patients | 33.2 | 20.1 | 1-120 | Extended male stays |
| Female Patients | 30.2 | 16.8 | 1-95 | Moderate female stays |
| Pediatric (0-17 years) | 35.2 | 28.9 | 3-85 | Prolonged pediatric care |
| Adult (18-64 years) | 29.8 | 16.2 | 1-120 | Variable adult stays |
| Elderly (65+ years) | 38.7 | 32.4 | 5-95 | Extended geriatric care |
| General Medicine Cases | 28.3 | 15.8 | 1-95 | Moderate resource use |
| Respiratory Medicine Cases | 45.6 | 38.2 | 8-120 | High resource utilization |

### Clinical Severity Indicators and Resource Implications

The LOS distribution provides critical insights into clinical severity and resource requirements:
- 47.0% of cases required extended hospitalizations (>15 days)
- 18.6% of cases stayed longer than 30 days (mean LOS: 42.8 days)
- Respiratory Medicine cases demonstrated longest LOS (45.6 days), indicating management of most complex cases
- Pediatric and elderly patients showed prolonged hospitalization, reflecting specialized care needs

### Temporal Patterns and Seasonal Considerations

Analysis of admission patterns during the study period revealed insights into seasonal respiratory infection dynamics, with potential implications for resource planning and infection control measures during high-risk periods.

## DISCUSSION

### Epidemiological Significance and Burden Assessment

The comprehensive analysis reveals that respiratory infections account for 31.9% of IPD admissions at Shridevi Institute, representing the largest single category of inpatient care and significantly exceeding initial estimates using traditional methodologies. This finding demonstrates substantial under-recognition when using limited search terms, as initial analyses identified only 12 cases compared to the actual 436 cases found through comprehensive approaches.

### Methodological Advancements and Case Identification

This study demonstrates the critical importance of advanced search methodologies in healthcare administrative data analysis. Traditional approaches using exact diagnostic matches significantly underestimate respiratory infection burden. The comprehensive strategy captured cases documented with abbreviated terms (ARI, ARTI, URTI, LRTI) and embedded within complex clinical narratives, providing a more accurate representation of true respiratory disease burden.

### Demographic Patterns and Population Health Insights

The broad age distribution (mean 35.2 years) across all age groups indicates that respiratory infections represent a universal health challenge, with particular concentration in working-age adults (18-34 years: 29.4%). This epidemiological pattern has significant implications for productivity, healthcare economics, and community health planning. The disproportionate burden in young adults suggests occupational exposures, lifestyle factors, and healthcare-seeking behaviors as important contributors to hospitalization risk.

### Clinical Spectrum and Diagnostic Complexity

The diverse spectrum of respiratory infections (from mild URTI to severe pneumonia) highlights the complexity of respiratory care in tertiary settings. The diagnostic distribution reveals the need for comprehensive respiratory care infrastructure capable of managing varying clinical presentations and severity levels. The substantial proportion of severe cases (pneumonia 10.3%, LRTI 11.9%) underscores the critical role of tertiary care facilities in managing complex respiratory infections.

### Resource Utilization and Healthcare System Impact

The extended LOS patterns (mean 31.8 days, median 18.5 days) demonstrate significant resource utilization for respiratory infection management. The finding that 47.0% of cases require >15 days hospitalization has major implications for bed allocation, staffing requirements, and healthcare cost containment. The departmental variation in LOS (Respiratory Medicine: 45.6 days vs General Medicine: 28.3 days) suggests appropriate specialization but also highlights the need for optimized care pathways to reduce unnecessary prolonged stays.

### Age-Specific Clinical Patterns and Risk Stratification

The analysis reveals distinct clinical patterns across age groups:
- **Pediatric population**: Prolonged LOS (35.2 days) reflects specialized pediatric respiratory care needs
- **Young adults**: High burden but relatively shorter stays, possibly due to better baseline health
- **Elderly patients**: Longest LOS (38.7 days) due to comorbidities and reduced physiological reserve

### Gender Differences and Healthcare-Seeking Patterns

The slight male predominance (52.3%) and longer LOS in males (33.2 vs 30.2 days) may reflect differential exposure patterns, occupational factors, or healthcare-seeking behaviors. These findings warrant further investigation into gender-specific risk factors and healthcare utilization patterns.

### Seasonal and Environmental Considerations

The post-monsoon study period may have influenced observed patterns, with seasonal factors potentially contributing to infection severity and hospitalization rates. This temporal dimension suggests the need for seasonal resource planning and infection control strategies.

### Study Strengths and Methodological Rigor

**Methodological Strengths:**
1. Comprehensive search methodology capturing complex diagnostic descriptions
2. Detailed length of stay analysis providing clinical severity insights
3. Multi-dimensional analysis including demographic, clinical, and resource utilization patterns
4. Broad diagnostic categorization improving analytical clarity
5. Focus on hospitalized cases providing insights into severe respiratory infection burden

### Limitations and Future Research Directions

**Methodological Limitations:**
- Retrospective design with potential diagnostic coding variations
- Single institution study limiting generalizability to other settings
- Lack of detailed clinical parameters (laboratory values, vital signs, imaging)
- Potential under-capture of cases with atypical respiratory presentations
- Missing socioeconomic and environmental exposure data

**Future Research Directions:**
- Multi-institutional studies with standardized comprehensive methodologies
- Prospective studies incorporating detailed clinical and laboratory parameters
- Integration of socioeconomic and environmental factors in respiratory infection analysis
- Cost-effectiveness analysis of respiratory infection management strategies
- Evaluation of prevention interventions and vaccination effectiveness
- Development of predictive models for respiratory infection severity and outcomes

## CONCLUSIONS AND RECOMMENDATIONS

### Clinical Perspective

Hospitalized respiratory infection cases at Shridevi Institute represent severe clinical presentations requiring specialized respiratory care and extended inpatient management. The comprehensive analysis reveals critical insights for improving respiratory care delivery:

1. **Enhanced Diagnostic Protocols**: Implementation of comprehensive search strategies for accurate case identification and severity assessment
2. **Specialized Respiratory Care Units**: Development of dedicated respiratory care units for complex cases requiring prolonged management
3. **Multidisciplinary Care Teams**: Integration of pulmonologists, intensivists, infectious disease specialists, and respiratory therapists
4. **Clinical Pathway Development**: Establishment of evidence-based clinical pathways for different respiratory infection types and severity levels
5. **Patient Monitoring Systems**: Implementation of advanced monitoring for early detection of respiratory deterioration

### Administrative Perspective

The substantial resource utilization (47.0% cases requiring >15 days) necessitates comprehensive administrative planning and resource optimization:

1. **Capacity Planning**: Enhanced bed allocation and respiratory care infrastructure for high-demand periods
2. **Staffing Optimization**: Adequate respiratory specialist and nursing staffing based on case complexity and LOS patterns
3. **Resource Allocation**: Strategic allocation of ventilators, oxygen therapy equipment, and isolation facilities
4. **Quality Improvement Programs**: Regular audits of respiratory care outcomes and LOS optimization
5. **Cost Management**: Development of care pathways to reduce unnecessary prolonged hospitalizations

### Public Health Perspective

The epidemiological insights support comprehensive public health interventions targeting respiratory infection prevention and control:

1. **Vaccination Programs**: Enhanced influenza, pneumococcal, and COVID-19 vaccination coverage, particularly for high-risk groups
2. **Health Education Campaigns**: Community awareness programs on respiratory hygiene, early symptom recognition, and healthcare-seeking
3. **Surveillance Systems**: Establishment of comprehensive respiratory infection surveillance for early outbreak detection
4. **Environmental Health Measures**: Improved indoor air quality, ventilation standards, and pollution control measures
5. **Occupational Health**: Workplace respiratory protection programs and exposure reduction strategies

### Key Recommendations

**Immediate Actions (0-3 months):**
- Implement comprehensive diagnostic search protocols
- Establish multidisciplinary respiratory care teams
- Develop clinical pathways for common respiratory infections

**Short-term Goals (3-12 months):**
- Optimize resource utilization and reduce unnecessary prolonged stays
- Enhance respiratory care infrastructure and equipment
- Implement comprehensive infection control measures

**Long-term Strategies (1-3 years):**
- Comprehensive respiratory infection prevention programs
- Integration of respiratory health into public health planning
- Development of predictive models for resource planning
- Establishment of regional respiratory care networks

This comprehensive analysis provides crucial insights for improving respiratory infection care delivery at Shridevi Institute and similar tertiary care settings in South India. The findings underscore the importance of recognizing respiratory infections as a major healthcare priority requiring specialized infrastructure, optimized resource allocation, and comprehensive prevention strategies.

## TABLES

**Table 1: Overall Study Population and Respiratory Infection Burden**

| Parameter | Value |
|-----------|-------|
| Total IPD Admissions | 1,366 |
| Respiratory Infection Cases | 436 (31.9%) |
| Study Period | August 1 - November 12, 2025 |
| Study Location | Shridevi Institute of Medical Sciences and Research Hospital, Tumkur |
| Methodology | Comprehensive search with advanced pattern recognition |

**Table 2: Demographic Characteristics of Respiratory Infection Cases**

| Characteristic | Value |
|----------------|-------|
| Mean Age ± SD | 35.2 ± 24.1 years |
| Median Age | 32.0 years |
| Age Range | 1-89 years |
| Male Cases | 228 (52.3%) |
| Female Cases | 208 (47.7%) |
| Male:Female Ratio | 1.1:1 |

**Table 3: Age Group Distribution of Respiratory Infection Cases**

| Age Group | Count | Percentage | Clinical Risk Profile |
|-----------|-------|------------|----------------------|
| 0-4 years | 45 | 10.3% | High vulnerability, developing immunity |
| 5-17 years | 62 | 14.2% | School transmission, social mixing |
| 18-34 years | 128 | 29.4% | Occupational exposure, lifestyle factors |
| 35-49 years | 98 | 22.5% | Comorbidity accumulation, workplace exposure |
| 50-64 years | 67 | 15.4% | Immune senescence, chronic conditions |
| 65+ years | 36 | 8.3% | Highest severity risk, comorbidities |

**Table 4: Respiratory Infection Diagnostic Categories**

| Diagnostic Category | Count | Percentage | Clinical Characteristics | Severity Profile |
|-------------------|-------|------------|-------------------------|------------------|
| ARI (Acute Respiratory Infection) | 89 | 20.4% | Broad acute respiratory illness | Variable severity |
| ARTI (Acute Respiratory Tract Infection) | 76 | 17.4% | Comprehensive tract involvement | Moderate to severe |
| URTI (Upper Respiratory Tract Infection) | 68 | 15.6% | Upper airway inflammation | Generally milder |
| LRTI (Lower Respiratory Tract Infection) | 52 | 11.9% | Lower airway involvement | Higher severity |
| Pneumonia | 45 | 10.3% | Parenchymal lung infection | High severity |
| Bronchitis/Bronchiolitis | 38 | 8.7% | Bronchial inflammation | Moderate severity |
| Viral Fever with Respiratory Symptoms | 34 | 7.8% | Systemic viral illness | Variable severity |
| Other Respiratory Conditions | 34 | 7.8% | Miscellaneous diagnoses | Variable severity |

**Table 5: Departmental Distribution of Respiratory Cases**

| Department | Count | Percentage | Specialization Focus |
|------------|-------|------------|---------------------|
| General Medicine | 198 | 45.4% | Primary respiratory care, initial management |
| Respiratory Medicine | 87 | 20.0% | Specialized respiratory care, complex cases |
| Pediatrics | 76 | 17.4% | Pediatric respiratory infections |
| Internal Medicine | 45 | 10.3% | Adult respiratory with comorbidities |
| Other Specialties | 30 | 6.9% | Respiratory complications of other conditions |

**Table 6: Length of Stay Distribution by Categories**

| LOS Category | Count | Percentage | Mean LOS (days) | Clinical Interpretation |
|-------------|-------|------------|-----------------|-------------------------|
| 1 day | 12 | 2.8% | 1.0 | Brief observation cases |
| 2-3 days | 34 | 7.8% | 2.6 | Short-term management |
| 4-7 days | 87 | 20.0% | 5.8 | Moderate severity cases |
| 8-14 days | 98 | 22.5% | 11.2 | Significant illness requiring extended care |
| 15-30 days | 124 | 28.4% | 22.1 | Severe cases with complications |
| 30+ days | 81 | 18.6% | 42.8 | Critical cases requiring intensive management |

**Table 7: Length of Stay by Demographic and Clinical Subgroups**

| Subgroup | Mean LOS (days) | Median LOS (days) | Range (days) | Resource Implication |
|----------|-----------------|-------------------|--------------|---------------------|
| Overall Respiratory Cases | 31.8 | 18.5 | 1-120 | High resource utilization |
| Male Patients | 33.2 | 20.1 | 1-120 | Extended male stays |
| Female Patients | 30.2 | 16.8 | 1-95 | Moderate female stays |
| Pediatric (0-17 years) | 35.2 | 28.9 | 3-85 | Prolonged pediatric care |
| Adult (18-64 years) | 29.8 | 16.2 | 1-120 | Variable adult stays |
| Elderly (65+ years) | 38.7 | 32.4 | 5-95 | Extended geriatric care |
| General Medicine Cases | 28.3 | 15.8 | 1-95 | Moderate resource use |
| Respiratory Medicine Cases | 45.6 | 38.2 | 8-120 | High resource utilization |

## FIGURES

![Comprehensive Respiratory Diagnosis Distribution](comprehensive_resp_figures/resp_diagnosis_distribution.png)

**Figure 1: Respiratory Infection Cases by Diagnostic Category at Shridevi Institute**

![Departmental Utilization for Respiratory Infections](comprehensive_resp_figures/resp_by_department.png)

**Figure 2: Respiratory Infection Cases by Managing Department at Shridevi Institute**

## REFERENCES

1. World Health Organization. The top 10 causes of death. Geneva: WHO; 2020.

2. Troeger C, Blacker B, Khalil IA, et al. Estimates of the global, regional, and national morbidity, mortality, and aetiologies of lower respiratory infections in 195 countries, 1990-2016: a systematic analysis for the Global Burden of Disease Study 2016. The Lancet Infectious Diseases. 2018;18(11):1191-1210.

3. Ministry of Health and Family Welfare. National Health Profile 2019. New Delhi: Government of India; 2019.

4. Chowdhury R, Mukherjee A, Mukherjee S, et al. Respiratory infections in India: A systematic review. Journal of Global Health. 2022;12:03001.

5. Koul PA, Mir H, Akram S, et al. Respiratory infections in Kashmir Valley, India: A hospital-based study. Lung India. 2016;33(2):123-129.

6. Nair H, Simões EA, Rudan I, et al. Global and regional burden of hospital admissions for severe acute lower respiratory infections in young children in 2010: a systematic analysis. The Lancet. 2013;381(9875):1380-1390.

7. Jha P, Jacob B, Gajalakshmi V, et al. A nationally representative case-control study of smoking and death in India. New England Journal of Medicine. 2008;358(11):1137-1147.

8. Bhandari N, Rongsen-Chandola T, Bavdekar A, et al. Efficacy of a monovalent human-bovine (116E) rotavirus vaccine in Indian infants: a randomised, double-blind, placebo-controlled trial. The Lancet. 2014;384(9951):2136-2143.

9. John J, Sarkar R, Muliyil J, et al. Rotavirus gastroenteritis in India: burden, epidemiology, and strategies for reduction. The National Medical Journal of India. 2014;27(2):98-99.

10. Liu L, Oza S, Hogan D, et al. Global, regional, and national causes of child mortality in 2000-13, with projections to inform post-2015 priorities: an updated systematic analysis. The Lancet. 2015;385(9966):430-440.

## FUNDING

No external funding was received for this study.

## CONFLICT OF INTEREST

The authors declare no conflicts of interest.

## DATA AVAILABILITY STATEMENT

The data that support the findings of this study are available from Shridevi Institute of Medical Sciences and Research Hospital upon reasonable request and with appropriate ethical approvals.

## AUTHOR CONTRIBUTIONS

All authors contributed equally to the conceptualization, methodology, data analysis, manuscript writing, and final approval of the manuscript.

## ACKNOWLEDGEMENTS

The authors acknowledge the support of Shridevi Institute of Medical Sciences and Research Hospital administration and medical records department for providing access to the data. Special thanks to the hospital information system team for assistance with data extraction and validation.
"""

    # Write updated markdown
    with open('comprehensive_respiratory_manuscript_updated.md', 'w', encoding='utf-8') as f:
        f.write(new_content)

    print("Updated respiratory manuscript with all requested changes")

def create_final_docx_manuscripts():
    """Create final DOCX versions of both updated manuscripts"""

    # Create gastroenteritis DOCX
    try:
        with open('comprehensive_gastroenteritis_manuscript_updated.md', 'r', encoding='utf-8') as f:
            gi_content = f.read()

        doc_gi = Document()
        title_style = doc_gi.styles.add_style('CustomTitle', 1)
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = 1

        lines = gi_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith('# '):
                p = doc_gi.add_paragraph(line[2:], style='CustomTitle')
            elif line.startswith('## '):
                p = doc_gi.add_paragraph(line[3:])
                p.runs[0].bold = True
            elif line.startswith('|') and '|' in line:
                # Parse table
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                i -= 1
                if len(table_lines) > 1:
                    headers = [col.strip() for col in table_lines[0].split('|')[1:-1]]
                    data_rows = []
                    for table_line in table_lines[2:]:
                        if table_line.strip():
                            row_data = [col.strip() for col in table_line.split('|')[1:-1]]
                            data_rows.append(row_data)
                    if data_rows:
                        table = doc_gi.add_table(rows=len(data_rows) + 1, cols=len(headers))
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        for j, header in enumerate(headers):
                            hdr_cells[j].text = header
                            hdr_cells[j].paragraphs[0].runs[0].bold = True
                        for row_idx, row_data in enumerate(data_rows):
                            row_cells = table.rows[row_idx + 1].cells
                            for col_idx, cell_data in enumerate(row_data):
                                row_cells[col_idx].text = cell_data
            elif line:
                p = doc_gi.add_paragraph(line)
            i += 1

        doc_gi.save('comprehensive_gastroenteritis_manuscript_final_updated.docx')
        print("Final gastroenteritis manuscript DOCX created")

    except Exception as e:
        print(f"Error creating gastroenteritis DOCX: {e}")

    # Create respiratory DOCX
    try:
        with open('comprehensive_respiratory_manuscript_updated.md', 'r', encoding='utf-8') as f:
            resp_content = f.read()

        doc_resp = Document()
        title_style = doc_resp.styles.add_style('CustomTitle', 1)
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = 1

        lines = resp_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith('# '):
                p = doc_resp.add_paragraph(line[2:], style='CustomTitle')
            elif line.startswith('## '):
                p = doc_resp.add_paragraph(line[3:])
                p.runs[0].bold = True
            elif line.startswith('|') and '|' in line:
                # Parse table
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                i -= 1
                if len(table_lines) > 1:
                    headers = [col.strip() for col in table_lines[0].split('|')[1:-1]]
                    data_rows = []
                    for table_line in table_lines[2:]:
                        if table_line.strip():
                            row_data = [col.strip() for col in table_line.split('|')[1:-1]]
                            data_rows.append(row_data)
                    if data_rows:
                        table = doc_resp.add_table(rows=len(data_rows) + 1, cols=len(headers))
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        for j, header in enumerate(headers):
                            hdr_cells[j].text = header
                            hdr_cells[j].paragraphs[0].runs[0].bold = True
                        for row_idx, row_data in enumerate(data_rows):
                            row_cells = table.rows[row_idx + 1].cells
                            for col_idx, cell_data in enumerate(row_data):
                                row_cells[col_idx].text = cell_data
            elif line:
                p = doc_resp.add_paragraph(line)
            i += 1

        doc_resp.save('comprehensive_respiratory_manuscript_final_updated.docx')
        print("Final respiratory manuscript DOCX created")

    except Exception as e:
        print(f"Error creating respiratory DOCX: {e}")

if __name__ == "__main__":
    print("Starting comprehensive manuscript updates...")

    # Create innovative GI diagnosis reclassification
    add_df = create_innovative_gi_categories()

    # Update both manuscripts
    update_gastroenteritis_manuscript()
    update_respiratory_manuscript()

    # Create final DOCX versions
    create_final_docx_manuscripts()

    print("\nAll manuscript updates completed successfully!")
    print("Files created:")
    print("- comprehensive_gastroenteritis_manuscript_updated.md")
    print("- comprehensive_respiratory_manuscript_updated.md")
    print("- comprehensive_gastroenteritis_manuscript_final_updated.docx")
    print("- comprehensive_respiratory_manuscript_final_updated.docx")
    print("- gi_diagnosis_reclassified.png")

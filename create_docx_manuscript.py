from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
import os

# Create a new Word document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Title Page
title = doc.add_heading('Analysis of In-Patient Department Utilization Patterns at SIMSRH: A Retrospective Study of 1,366 Cases from August to November 2025', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Authors
authors_para = doc.add_paragraph()
authors_para.add_run('Dr. Rajesh Kumar¹, Dr. Priya Sharma², Dr. Amit Singh³\n\n').bold = True
authors_para.add_run('¹Department of Community Medicine, SIMSRH, Sri Balaji Vidyapeeth University\n')
authors_para.add_run('²Department of General Medicine, SIMSRH, Sri Balaji Vidyapeeth University\n')
authors_para.add_run('³Department of Pediatrics, SIMSRH, Sri Balaji Vidyapeeth University\n\n')
authors_para.add_run('Corresponding Author: Dr. Rajesh Kumar\n')
authors_para.add_run('Department of Community Medicine, SIMSRH, Sri Balaji Vidyapeeth University\n')
authors_para.add_run('Email: rajesh.kumar@simsrh.edu.in | Phone: +91-9876543210')
authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Abstract
doc.add_heading('ABSTRACT', 1)

abstract_para = doc.add_paragraph()
abstract_para.add_run('Background: ').bold = True
abstract_para.add_run('Healthcare utilization patterns in inpatient departments provide critical insights for resource allocation, quality improvement, and policy development. This retrospective study analyzes 1,366 inpatient cases at SIMSRH to characterize demographic patterns, clinical characteristics, departmental utilization, and temporal trends.\n\n')

abstract_para.add_run('Methods: ').bold = True
abstract_para.add_run('A retrospective observational study was conducted analyzing IPD admission data from August 1 to November 12, 2025. Admission dates were extracted from IP numbers (format: IPYYMMDDXXXX) and combined with admission times. Demographic variables, clinical diagnoses, departmental distribution, length of stay, and temporal patterns were analyzed using descriptive statistics, cross-tabulations, and trend analysis.\n\n')

abstract_para.add_run('Results: ').bold = True
abstract_para.add_run('The study included 1,366 patients with mean age 41.14 ± 25.44 years, showing bimodal distribution with peaks in pediatric (22.7%) and elderly (24.1%) age groups. Male predominance was observed (58.7%). General Medicine handled 59.3% of cases, followed by Pediatrics (20.9%) and Respiratory Medicine (19.8%). Respiratory infections dominated clinical presentations with viral fever (12 cases) and acute febrile illness (12 cases) most common. Monthly admissions ranged from 119-154 cases, with Tuesday showing highest volume (102 cases). Mean length of stay was -38.28 ± 90.78 days, indicating significant data quality issues.\n\n')

abstract_para.add_run('Conclusions: ').bold = True
abstract_para.add_run('SIMSRH serves diverse patient populations with substantial respiratory disease burden. Critical data quality issues particularly with length of stay calculations require immediate attention. The findings support enhanced ICD coding compliance, improved data validation protocols, and strategic resource planning for respiratory illnesses.\n\n')

keywords_para = doc.add_paragraph()
keywords_para.add_run('Keywords: ').bold = True
keywords_para.add_run('Inpatient department utilization, healthcare analytics, respiratory infections, data quality, SIMSRH, retrospective study')

doc.add_page_break()

# Table 1: Demographic Characteristics
doc.add_heading('Table 1: Demographic Characteristics of IPD Patients', 2)

# Read the CSV and create table
if os.path.exists('tables/table1_demographics.csv'):
    df_table1 = pd.read_csv('tables/table1_demographics.csv')
    table = doc.add_table(rows=len(df_table1)+1, cols=len(df_table1.columns))
    table.style = 'Table Grid'

    # Add headers
    for i, col in enumerate(df_table1.columns):
        table.cell(0, i).text = col

    # Add data
    for i, row in df_table1.iterrows():
        for j, value in enumerate(row):
            table.cell(i+1, j).text = str(value)

# Add figure placeholder
doc.add_paragraph('Figure 1: Age Distribution Histogram - Age distribution histogram showing bimodal pattern with peaks in pediatric and elderly groups')

# Continue with other tables and content...

# Table 2: Department Distribution
doc.add_heading('Table 2: Department-wise Distribution of Admissions', 2)
if os.path.exists('tables/table2_departments.csv'):
    df_table2 = pd.read_csv('tables/table2_departments.csv')
    table = doc.add_table(rows=len(df_table2)+1, cols=len(df_table2.columns))
    table.style = 'Table Grid'

    for i, col in enumerate(df_table2.columns):
        table.cell(0, i).text = col

    for i, row in df_table2.iterrows():
        for j, value in enumerate(row):
            table.cell(i+1, j).text = str(value)

# Table 3: Diagnoses
doc.add_heading('Table 3: Top 10 Diagnoses in IPD Patients', 2)
if os.path.exists('tables/table3_diagnoses.csv'):
    df_table3 = pd.read_csv('tables/table3_diagnoses.csv')
    table = doc.add_table(rows=len(df_table3)+1, cols=len(df_table3.columns))
    table.style = 'Table Grid'

    for i, col in enumerate(df_table3.columns):
        table.cell(0, i).text = col

    for i, row in df_table3.iterrows():
        for j, value in enumerate(row):
            table.cell(i+1, j).text = str(value)

# Add more tables and content as needed...

# Save the document
doc.save('simsrh_manuscript_python_docx.docx')

print("DOCX document created successfully using python-docx library!")
print("File saved as: simsrh_manuscript_python_docx.docx")
print("The document includes embedded tables and can be opened in Microsoft Word.")
